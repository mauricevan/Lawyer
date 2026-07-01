"""Export conversations to PDF and Word."""
import io
from datetime import datetime

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession

from backend.src.database import SessionLocal
from backend.src.services.conversation_service import ConversationService

router = APIRouter(prefix="/export", tags=["export"])
service = ConversationService()


async def get_db() -> AsyncSession:
    async with SessionLocal() as session:
        yield session


@router.get("/pdf/{conversation_id}")
async def export_pdf(
    conversation_id: str,
    session: AsyncSession = Depends(get_db),
) -> StreamingResponse:
    content = await _build_export_content(session, conversation_id)
    pdf_bytes = _generate_pdf(content)
    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=gesprek-{conversation_id[:8]}.pdf"},
    )


@router.get("/docx/{conversation_id}")
async def export_docx(
    conversation_id: str,
    session: AsyncSession = Depends(get_db),
) -> StreamingResponse:
    content = await _build_export_content(session, conversation_id)
    docx_bytes = _generate_docx(content)
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=gesprek-{conversation_id[:8]}.docx"},
    )


async def _build_export_content(session: AsyncSession, conversation_id: str) -> dict:
    conv = await service.get_conversation(session, conversation_id)
    if not conv:
        raise HTTPException(status_code=404, detail="Gesprek niet gevonden")
    messages = await service.list_messages(session, conversation_id)
    return {
        "title": conv.title,
        "exported_at": datetime.utcnow().isoformat(),
        "messages": [
            {"role": m.role, "content": m.content, "citations": m.citations or []}
            for m in messages
        ],
    }


def _generate_pdf(content: dict) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    y = 800
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, content["title"][:80])
    y -= 30
    c.setFont("Helvetica", 10)
    for msg in content["messages"]:
        role = "Vraag" if msg["role"] == "user" else "Antwoord"
        c.drawString(50, y, f"{role}:")
        y -= 15
        for line in msg["content"][:500].split("\n"):
            c.drawString(60, y, line[:90])
            y -= 12
            if y < 50:
                c.showPage()
                y = 800
        for cite in msg.get("citations", []):
            y -= 12
            legal = cite.get("legal_citation", cite.get("celex", ""))
            c.drawString(60, y, f"Bron: {legal}"[:90])
            y -= 12
        y -= 20
    c.save()
    return buffer.getvalue()


def _generate_docx(content: dict) -> bytes:
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_heading(content["title"], 0)
    doc.add_paragraph(f"Geëxporteerd: {content['exported_at']}")
    for msg in content["messages"]:
        role = "Vraag" if msg["role"] == "user" else "Antwoord"
        doc.add_heading(role, level=2)
        doc.add_paragraph(msg["content"])
        for cite in msg.get("citations", []):
            legal = cite.get("legal_citation", cite.get("celex", ""))
            doc.add_paragraph(f"Bron: {legal}", style="Intense Quote")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
